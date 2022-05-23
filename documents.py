from docx import Document
from datetime import datetime
from mailer import send_mail_with_attachment
from openpyxl import load_workbook

from googletrans import Translator


def modifyDoc(document, key, msg):
    for paragraph in document.paragraphs:
        if key in paragraph.text:
            paragraph.text = msg


def modifyTableDoc(document, key, msg):
    for table in document.tables:
        for i, row in enumerate(table.rows):
            for cell in row.cells:
                if key in cell.text:
                    cell.text = msg


def fill_test_plan_desc(spreadsheet, desc):
    finished = True
    i = 4
    while (finished):
        if spreadsheet["A" + str(i)].value is None:
            spreadsheet["A" + str(i)] = "TSC-" + str(i - 3)
            spreadsheet["B" + str(i)] = desc
            finished = False
        else:
            i = i + 1


def fill_test_plan_status(spreadsheet, status):
    finished = True
    i = 4
    while (finished):
        if spreadsheet["A" + str(i + 1)].value is None:
            spreadsheet["C" + str(i)] = status
            finished = False
        else:
            i = i + 1


def addTitle(document, title):
    document.add_heading(title, 2)


def addDescription(document, desc):
    document.add_paragraph(desc)


def create_psd(msg, question, mail):
    filename = mail + "PSD.docx"
    if "PSD in creation, what do you want as a title?" == question:
        document = Document(filename)
        modifyDoc(document, "xtitlex", msg)
        today = datetime.today().strftime('%Y-%m-%d')
        modifyDoc(document, "xdatex", today)
        modifyTableDoc(document, "xdatex", today)
        modifyTableDoc(document, "xauthorx", mail)
        document.save(filename)
        return "what is the client context?"
    elif "what is the client context?" == question:
        document = Document(filename)
        modifyDoc(document, "xclient contextx", msg)
        document.save(filename)
        return "what is the business context?"
    elif "what is the business context?" == question:
        document = Document(filename)
        modifyDoc(document, "xbusiness contextx", msg)
        document.save(filename)
        return "give me a brief description of the change request ?"
    elif "give me a brief description of the change request ?" == question:
        document = Document(filename)
        modifyDoc(document, "Xdescriptionx", msg)
        document.save(filename)
        return "do you want to add new feature description ?"
    elif "do you want to add new feature description ?" == question and ("no" not in msg):
        return "what is the feature's title ?"
    elif "what is the feature's title ?" == question:
        document = Document(filename)
        addTitle(document, msg)
        document.save(filename)
        return "what is the description for this feature ?"
    elif "what is the description for this feature ?" == question:
        document = Document(filename)
        addDescription(document, msg)
        document.save(filename)
        return "do you want to add new feature description ?"
    else:
        send_mail_with_attachment(filename, mail, "Generated PSD File")
        return "PSD is now saved and is gonna be sent to you on mail."


def create_pfr(msg, question, mail):
    filename = mail + "PFR.docx"
    if "PFR in creation, what do you want as a title?" == question:
        document = Document(filename)
        modifyDoc(document, "xtitlex", msg)
        today = datetime.today().strftime('%Y-%m-%d')
        modifyDoc(document, "xdatex", today)
        modifyTableDoc(document, "xdatex", today)
        modifyTableDoc(document, "xauthorx", mail)
        document.save(filename)
        return "what is the aim of the document?"
    elif "what is the aim of the document?" == question:
        document = Document(filename)
        modifyDoc(document, "xaimx", msg)
        document.save(filename)
        return "describe the current behavior?"
    elif "describe the current behavior?" == question:
        document = Document(filename)
        modifyDoc(document, "xcurrent behaviorx", msg)
        document.save(filename)
        return "what is the proposed solution?"
    elif "what is the proposed solution?" == question:
        document = Document(filename)
        modifyDoc(document, "xsolutionx", msg)
        document.save(filename)
        return "do you want to add another new feature ?"
    elif "do you want to add another new feature ?" == question and ("no" not in msg):
        return "what is the new feature's title ?"
    elif "what is the new feature's title ?" == question:
        document = Document(filename)
        addTitle(document, msg)
        document.save(filename)
        return "what is the description ?"
    elif "what is the description ?" == question:
        document = Document(filename)
        addDescription(document, msg)
        document.save(filename)
        return "do you want to add another new feature ?"
    else:
        send_mail_with_attachment(filename, mail, "Generated PFR File")
        return "PFR is now saved and is gonna be sent to you on mail."


def create_test_plan(msg, question, mail):
    workbook_path = mail + "plan.xlsx"
    workbook = load_workbook(filename=workbook_path)
    spreadsheet = workbook.active

    if "test plan in creation, what is the title of the test?" == question:
        spreadsheet["B1"] = msg
        workbook.save(workbook_path)
        return "what is the description of the test case?"
    elif "what is the description of the test case?" == question:
        fill_test_plan_desc(spreadsheet, msg)
        workbook.save(workbook_path)
        return "what is the actual status of this test case ?"
    elif "what is the actual status of this test case ?" == question:
        fill_test_plan_status(spreadsheet, msg)
        workbook.save(workbook_path)
        return "do you want to add another test case ?"
    elif "do you want to add another test case ?" == question and ("no" not in msg):
        return "what is the description of the test case?"
    else:
        send_mail_with_attachment(workbook_path, mail, "Generated Test Plan File")
        return "test plan is now saved and is gonna be sent to you on mail."

    ################## Fr


def traduireEnAnglais(text):
    if text != "":
        translator = Translator()
        translation = translator.translate(text, dest="en")
        return translation.text


def create_psd_fr(message, question, mail):
    filename = mail + "PSD.docx"
    msg = traduireEnAnglais(message)
    if "PSD en création, tu veux quoi comme titre ?" == question:
        document = Document(filename)
        modifyDoc(document, "xtitlex", msg)
        today = datetime.today().strftime('%Y-%m-%d')
        modifyDoc(document, "xdatex", today)
        modifyTableDoc(document, "xdatex", today)
        modifyTableDoc(document, "xauthorx", mail)
        document.save(filename)
        return "quel est le contexte client ?"
    elif "quel est le contexte client ?" == question:
        document = Document(filename)
        modifyDoc(document, "xclient contextx", msg)
        document.save(filename)
        return "quel est le contexte commercial ?"
    elif "quel est le contexte commercial ?" == question:
        document = Document(filename)
        modifyDoc(document, "xbusiness contextx", msg)
        document.save(filename)
        return "me donner une brève description de la demande de changement ?"
    elif "me donner une brève description de la demande de changement ?" == question:
        document = Document(filename)
        modifyDoc(document, "Xdescriptionx", msg)
        document.save(filename)
        return "voulez-vous ajouter une nouvelle fonctionnalité?"
    elif "voulez-vous ajouter une nouvelle fonctionnalité?" == question and ("non" not in message):
        return "quel est le titre de la fonctionnalité ?"
    elif "quel est le titre de la fonctionnalité ?" == question:
        document = Document(filename)
        addTitle(document, msg)
        document.save(filename)
        return "quelle est la description de cette fonctionnalité ?"
    elif "quelle est la description de cette fonctionnalité ?" == question:
        document = Document(filename)
        addDescription(document, msg)
        document.save(filename)
        return "voulez-vous ajouter une nouvelle fonctionnalité?"
    else:
        send_mail_with_attachment(filename, mail, "Fichier PSD généré")
        return "PSD est maintenant enregistré et va vous être envoyé par courrier."


def create_pfr_fr(message, question, mail):
    filename = mail + "PFR.docx"
    msg = traduireEnAnglais(message)
    if "PFR en création, tu veux quoi comme titre ?" == question:
        document = Document(filename)
        modifyDoc(document, "xtitlex", msg)
        today = datetime.today().strftime('%Y-%m-%d')
        modifyDoc(document, "xdatex", today)
        modifyTableDoc(document, "xdatex", today)
        modifyTableDoc(document, "xauthorx", mail)
        document.save(filename)
        return "quel est l'objectif du document ?"
    elif "quel est l'objectif du document ?" == question:
        document = Document(filename)
        modifyDoc(document, "xaimx", msg)
        document.save(filename)
        return "décrire l'existant ?"
    elif "décrire l'existant ?" == question:
        document = Document(filename)
        modifyDoc(document, "xcurrent behaviorx", msg)
        document.save(filename)
        return "quelle est la solution proposée ?"
    elif "quelle est la solution proposée ?" == question:
        document = Document(filename)
        modifyDoc(document, "xsolutionx", msg)
        document.save(filename)
        return "voulez-vous ajouter une autre nouvelle fonctionnalité ?"
    elif "voulez-vous ajouter une autre nouvelle fonctionnalité ?" == question and ("non" not in message):
        return "quel est le titre de la nouvelle fonctionnalité?"
    elif "quel est le titre de la nouvelle fonctionnalité?" == question:
        document = Document(filename)
        addTitle(document, msg)
        document.save(filename)
        return "c'est quoi le descriptif ?"
    elif "c'est quoi le descriptif ?" == question:
        document = Document(filename)
        addDescription(document, msg)
        document.save(filename)
        return "voulez-vous ajouter une autre nouvelle fonctionnalité ?"
    else:
        send_mail_with_attachment(filename, mail, "Fichier PFR généré")
        return "PFR est maintenant enregistré et va vous être envoyé par courrier."


def create_test_plan_fr(message, question, mail):
    workbook_path = mail + "plan.xlsx"
    msg = traduireEnAnglais(message)
    workbook = load_workbook(filename=workbook_path)
    spreadsheet = workbook.active

    if "plan de test en création, quel est le titre du test ?" == question:
        spreadsheet["B1"] = msg
        workbook.save(workbook_path)
        return "quelle est la description du cas de test ?"
    elif "quelle est la description du cas de test ?" == question:
        fill_test_plan_desc(spreadsheet, msg)
        workbook.save(workbook_path)
        return "quel est le statut actuel de ce cas de test ?"
    elif "quel est le statut actuel de ce cas de test ?" == question:
        fill_test_plan_status(spreadsheet, msg)
        workbook.save(workbook_path)
        return "voulez-vous ajouter un autre cas de test ?"
    elif "voulez-vous ajouter un autre cas de test ?" == question and ("non" not in message):
        return "quelle est la description du cas de test ?"
    else:
        send_mail_with_attachment(workbook_path, mail, "Fichier de plan de test généré")
        return "Le plan de test est maintenant enregistré et va vous être envoyé par courrier."
